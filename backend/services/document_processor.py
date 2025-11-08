"""
Production-grade document processing service for Indian Legal AI Assistant
Handles parsing, chunking, and indexing of legal documents with high precision
"""
import asyncio
import hashlib
import uuid
from typing import List, Dict, Any, Optional, Tuple, BinaryIO
from pathlib import Path
from datetime import datetime
import structlog
from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession

# Document parsing libraries
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pdfplumber
from azure.ai.formrecognizer.aio import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential

from config import settings
from database import Document, DocumentChunk, DocumentType, ProcessingStatus
from core.vector_search import VectorSearchEngine
from services.azure_storage import AzureStorageService
from utils.text_processing import TextProcessor, LegalTextChunker
from utils.metadata_extractor import LegalMetadataExtractor

logger = structlog.get_logger(__name__)

class DocumentProcessor:
    """
    Production-grade document processing service
    
    Features:
    - Multi-format document parsing (PDF, DOCX, TXT)
    - OCR for scanned documents
    - Legal-specific text chunking
    - Metadata extraction and classification
    - Asynchronous processing pipeline
    - Error recovery and retry logic
    """
    
    def __init__(self):
        self.vector_search = VectorSearchEngine()
        self.storage_service = AzureStorageService()
        self.text_processor = TextProcessor()
        self.chunker = LegalTextChunker()
        self.metadata_extractor = LegalMetadataExtractor()
        
        # Initialize Azure Form Recognizer for OCR
        if settings.azure_openai.api_key:  # Using same credentials for simplicity
            self.form_recognizer = DocumentAnalysisClient(
                endpoint=settings.azure_openai.endpoint.replace('openai', 'cognitiveservices'),
                credential=AzureKeyCredential(settings.azure_openai.api_key)
            )
        else:
            self.form_recognizer = None
        
        # Processing statistics
        self.processing_stats = {
            'documents_processed': 0,
            'total_chunks_created': 0,
            'processing_errors': 0,
            'ocr_operations': 0
        }
        
        logger.info("Document processor initialized with advanced parsing capabilities")
    
    async def process_upload(
        self,
        file: UploadFile,
        user_id: str,
        session: AsyncSession,
        document_type: Optional[DocumentType] = None
    ) -> str:
        """
        Process uploaded document through complete pipeline
        
        Args:
            file: Uploaded file
            user_id: User ID for access control
            session: Database session
            document_type: Optional document type classification
            
        Returns:
            Document ID for tracking
        """
        document_id = str(uuid.uuid4())
        
        try:
            # Read file content
            file_content = await file.read()
            file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Check for duplicate documents
            existing_doc = await self._check_duplicate_document(file_hash, user_id, session)
            if existing_doc:
                logger.info("Duplicate document detected", document_id=existing_doc.id)
                return str(existing_doc.id)
            
            # Create document record
            document = await self._create_document_record(
                document_id=document_id,
                file=file,
                file_content=file_content,
                file_hash=file_hash,
                user_id=user_id,
                document_type=document_type,
                session=session
            )
            
            # Upload to Azure Storage
            blob_url = await self.storage_service.upload_document(
                document_id=document_id,
                filename=file.filename,
                content=file_content,
                user_id=user_id
            )
            
            # Update document with storage info
            document.blob_url = blob_url
            document.blob_container = settings.azure_storage.container_name
            document.blob_path = f"{user_id}/{document_id}/{file.filename}"
            
            await session.commit()
            
            # Start asynchronous processing
            asyncio.create_task(
                self._process_document_async(document_id, file_content, file.filename, user_id)
            )
            
            logger.info(
                "Document upload initiated",
                document_id=document_id,
                filename=file.filename,
                size_bytes=len(file_content)
            )
            
            return document_id
            
        except Exception as e:
            logger.error(
                "Document upload failed",
                document_id=document_id,
                filename=file.filename,
                error=str(e),
                exc_info=True
            )
            
            # Update document status to failed
            await self._update_document_status(
                document_id, ProcessingStatus.FAILED, 0, str(e), session
            )
            raise
    
    async def _process_document_async(
        self,
        document_id: str,
        file_content: bytes,
        filename: str,
        user_id: str
    ):
        """
        Asynchronous document processing pipeline
        """
        session = None
        
        try:
            from database.connection import get_async_session
            session = await get_async_session()
            
            # Update status to processing
            await self._update_document_status(
                document_id, ProcessingStatus.PROCESSING, 10, "Starting document processing", session
            )
            
            # Step 1: Extract text content
            logger.info("Starting text extraction", document_id=document_id)
            text_content, page_count = await self._extract_text_content(
                file_content, filename, document_id
            )
            
            await self._update_document_status(
                document_id, ProcessingStatus.PROCESSING, 30, "Text extraction completed", session
            )
            
            # Step 2: Extract metadata and classify document
            logger.info("Extracting metadata", document_id=document_id)
            metadata = await self.metadata_extractor.extract_metadata(
                text_content, filename
            )
            
            # Update document with extracted metadata
            await self._update_document_metadata(
                document_id, metadata, page_count, len(text_content.split()), session
            )
            
            await self._update_document_status(
                document_id, ProcessingStatus.PROCESSING, 50, "Metadata extraction completed", session
            )
            
            # Step 3: Chunk the document
            logger.info("Chunking document", document_id=document_id)
            chunks = await self.chunker.chunk_document(
                text=text_content,
                document_type=metadata.get('document_type', DocumentType.USER_DOCUMENT),
                source_name=filename,
                metadata=metadata
            )
            
            await self._update_document_status(
                document_id, ProcessingStatus.PROCESSING, 70, "Document chunking completed", session
            )
            
            # Step 4: Generate embeddings and store chunks
            logger.info("Generating embeddings", document_id=document_id)
            await self._process_and_store_chunks(
                chunks, document_id, user_id, session
            )
            
            await self._update_document_status(
                document_id, ProcessingStatus.PROCESSING, 90, "Embeddings generated", session
            )
            
            # Step 5: Final processing and indexing
            await self._finalize_document_processing(document_id, session)
            
            await self._update_document_status(
                document_id, ProcessingStatus.COMPLETED, 100, "Document processing completed successfully", session
            )
            
            # Update statistics
            self.processing_stats['documents_processed'] += 1
            self.processing_stats['total_chunks_created'] += len(chunks)
            
            logger.info(
                "Document processing completed successfully",
                document_id=document_id,
                chunks_created=len(chunks),
                processing_time="async"
            )
            
        except Exception as e:
            logger.error(
                "Async document processing failed",
                document_id=document_id,
                error=str(e),
                exc_info=True
            )
            
            self.processing_stats['processing_errors'] += 1
            
            if session:
                await self._update_document_status(
                    document_id, ProcessingStatus.FAILED, 0, f"Processing failed: {str(e)}", session
                )
        
        finally:
            if session:
                await session.close()
    
    async def _extract_text_content(
        self,
        file_content: bytes,
        filename: str,
        document_id: str
    ) -> Tuple[str, Optional[int]]:
        """
        Extract text content from various file formats
        """
        file_extension = Path(filename).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return await self._extract_pdf_text(file_content, document_id)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_docx_text(file_content)
            elif file_extension == '.txt':
                return await self._extract_txt_text(file_content)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(
                "Text extraction failed",
                document_id=document_id,
                file_extension=file_extension,
                error=str(e)
            )
            raise
    
    async def _extract_pdf_text(self, file_content: bytes, document_id: str) -> Tuple[str, int]:
        """
        Extract text from PDF with OCR fallback for scanned documents
        """
        text_parts = []
        page_count = 0
        
        try:
            # Try PyMuPDF first (faster for text-based PDFs)
            doc = fitz.open(stream=file_content, filetype="pdf")
            page_count = len(doc)
            
            for page_num in range(page_count):
                page = doc[page_num]
                text = page.get_text()
                
                # If page has very little text, it might be scanned
                if len(text.strip()) < 50:
                    # Try OCR with Azure Form Recognizer
                    ocr_text = await self._perform_ocr_on_page(page, document_id, page_num)
                    if ocr_text:
                        text_parts.append(ocr_text)
                        self.processing_stats['ocr_operations'] += 1
                    else:
                        text_parts.append(text)
                else:
                    text_parts.append(text)
            
            doc.close()
            
            # If very little text extracted, try full OCR
            full_text = "\n\n".join(text_parts)
            if len(full_text.strip()) < 100:
                logger.info("Attempting full document OCR", document_id=document_id)
                ocr_text = await self._perform_full_document_ocr(file_content)
                if ocr_text and len(ocr_text) > len(full_text):
                    full_text = ocr_text
                    self.processing_stats['ocr_operations'] += 1
            
            return full_text, page_count
            
        except Exception as e:
            logger.warning("PyMuPDF extraction failed, trying pdfplumber", error=str(e))
            
            # Fallback to pdfplumber
            try:
                import io
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                
                return "\n\n".join(text_parts), page_count
                
            except Exception as e2:
                logger.error("All PDF extraction methods failed", error=str(e2))
                raise
    
    async def _extract_docx_text(self, file_content: bytes) -> Tuple[str, Optional[int]]:
        """Extract text from DOCX files"""
        try:
            import io
            doc = DocxDocument(io.BytesIO(file_content))
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts), None
            
        except Exception as e:
            logger.error("DOCX extraction failed", error=str(e))
            raise
    
    async def _extract_txt_text(self, file_content: bytes) -> Tuple[str, Optional[int]]:
        """Extract text from TXT files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text, None
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use error handling
            text = file_content.decode('utf-8', errors='replace')
            return text, None
            
        except Exception as e:
            logger.error("TXT extraction failed", error=str(e))
            raise
    
    async def _perform_ocr_on_page(
        self,
        page: fitz.Page,
        document_id: str,
        page_num: int
    ) -> Optional[str]:
        """
        Perform OCR on a single PDF page using Azure Form Recognizer
        """
        if not self.form_recognizer:
            return None
        
        try:
            # Convert page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
            img_data = pix.tobytes("png")
            
            # Perform OCR
            async with self.form_recognizer:
                poller = await self.form_recognizer.begin_analyze_document(
                    "prebuilt-read", img_data
                )
                result = await poller.result()
            
            # Extract text from OCR result
            text_parts = []
            for page_result in result.pages:
                for line in page_result.lines:
                    text_parts.append(line.content)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.warning(
                "OCR failed for page",
                document_id=document_id,
                page_num=page_num,
                error=str(e)
            )
            return None
    
    async def _perform_full_document_ocr(self, file_content: bytes) -> Optional[str]:
        """
        Perform OCR on entire document using Azure Form Recognizer
        """
        if not self.form_recognizer:
            return None
        
        try:
            async with self.form_recognizer:
                poller = await self.form_recognizer.begin_analyze_document(
                    "prebuilt-read", file_content
                )
                result = await poller.result()
            
            # Extract all text
            text_parts = []
            for page in result.pages:
                page_text = []
                for line in page.lines:
                    page_text.append(line.content)
                text_parts.append("\n".join(page_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.warning("Full document OCR failed", error=str(e))
            return None
    
    async def _process_and_store_chunks(
        self,
        chunks: List[Dict[str, Any]],
        document_id: str,
        user_id: str,
        session: AsyncSession
    ):
        """
        Generate embeddings for chunks and store in database
        """
        from openai import AsyncAzureOpenAI
        
        client = AsyncAzureOpenAI(
            api_key=settings.azure_openai.api_key,
            api_version=settings.azure_openai.api_version,
            azure_endpoint=settings.azure_openai.endpoint
        )
        
        # Process chunks in batches for efficiency
        batch_size = 20
        
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i + batch_size]
            
            # Generate embeddings for batch
            texts = [chunk['text'] for chunk in batch]
            
            try:
                response = await client.embeddings.create(
                    input=texts,
                    model=settings.azure_openai.embedding_deployment
                )
                
                # Create DocumentChunk objects
                chunk_objects = []
                for j, (chunk, embedding_data) in enumerate(zip(batch, response.data)):
                    chunk_obj = DocumentChunk(
                        id=uuid.uuid4(),
                        document_id=document_id,
                        user_id=user_id,
                        text=chunk['text'],
                        text_hash=hashlib.sha256(chunk['text'].encode()).hexdigest(),
                        embedding=embedding_data.embedding,
                        chunk_index=i + j,
                        start_page=chunk.get('start_page'),
                        end_page=chunk.get('end_page'),
                        start_char=chunk.get('start_char'),
                        end_char=chunk.get('end_char'),
                        section_title=chunk.get('section_title'),
                        section_number=chunk.get('section_number'),
                        paragraph_number=chunk.get('paragraph_number'),
                        content_type=chunk.get('content_type'),
                        importance_score=chunk.get('importance_score'),
                        source_name=chunk['source_name'],
                        source_type=chunk['source_type'],
                        metadata=chunk.get('metadata', {})
                    )
                    chunk_objects.append(chunk_obj)
                
                # Bulk insert chunks
                session.add_all(chunk_objects)
                await session.commit()
                
                logger.debug(
                    "Batch of chunks processed",
                    document_id=document_id,
                    batch_size=len(batch),
                    batch_number=i // batch_size + 1
                )
                
            except Exception as e:
                logger.error(
                    "Chunk processing failed",
                    document_id=document_id,
                    batch_number=i // batch_size + 1,
                    error=str(e)
                )
                raise
    
    async def _create_document_record(
        self,
        document_id: str,
        file: UploadFile,
        file_content: bytes,
        file_hash: str,
        user_id: str,
        document_type: Optional[DocumentType],
        session: AsyncSession
    ) -> Document:
        """Create initial document record in database"""
        
        document = Document(
            id=document_id,
            user_id=user_id,
            filename=file.filename,
            original_filename=file.filename,
            file_size_bytes=len(file_content),
            file_hash=file_hash,
            mime_type=file.content_type or 'application/octet-stream',
            document_type=document_type or DocumentType.USER_DOCUMENT,
            status=ProcessingStatus.PENDING,
            processing_progress=0,
            processing_message="Document uploaded, processing queued",
            processing_started_at=datetime.utcnow(),
            language='en',  # Default to English, can be detected later
            is_public=False,
            access_level='private'
        )
        
        session.add(document)
        await session.commit()
        
        return document
    
    async def _update_document_status(
        self,
        document_id: str,
        status: ProcessingStatus,
        progress: int,
        message: str,
        session: AsyncSession
    ):
        """Update document processing status"""
        try:
            from sqlalchemy import select, update
            
            stmt = update(Document).where(
                Document.id == document_id
            ).values(
                status=status,
                processing_progress=progress,
                processing_message=message,
                processing_completed_at=datetime.utcnow() if status == ProcessingStatus.COMPLETED else None
            )
            
            await session.execute(stmt)
            await session.commit()
            
        except Exception as e:
            logger.error(
                "Failed to update document status",
                document_id=document_id,
                status=status.value,
                error=str(e)
            )
    
    async def _update_document_metadata(
        self,
        document_id: str,
        metadata: Dict[str, Any],
        page_count: Optional[int],
        word_count: int,
        session: AsyncSession
    ):
        """Update document with extracted metadata"""
        try:
            from sqlalchemy import update
            
            update_values = {
                'legal_metadata': metadata,
                'word_count': word_count
            }
            
            if page_count:
                update_values['page_count'] = page_count
            
            # Extract specific metadata fields
            if 'title' in metadata:
                update_values['title'] = metadata['title']
            if 'author' in metadata:
                update_values['author'] = metadata['author']
            if 'court' in metadata:
                update_values['court'] = metadata['court']
            if 'year' in metadata:
                update_values['year'] = metadata['year']
            if 'citation' in metadata:
                update_values['citation'] = metadata['citation']
            
            stmt = update(Document).where(
                Document.id == document_id
            ).values(**update_values)
            
            await session.execute(stmt)
            await session.commit()
            
        except Exception as e:
            logger.error(
                "Failed to update document metadata",
                document_id=document_id,
                error=str(e)
            )
    
    async def _check_duplicate_document(
        self,
        file_hash: str,
        user_id: str,
        session: AsyncSession
    ) -> Optional[Document]:
        """Check if document with same hash already exists for user"""
        try:
            from sqlalchemy import select
            
            stmt = select(Document).where(
                Document.file_hash == file_hash,
                Document.user_id == user_id,
                Document.status == ProcessingStatus.COMPLETED
            )
            
            result = await session.execute(stmt)
            return result.scalar_one_or_none()
            
        except Exception as e:
            logger.error("Duplicate check failed", error=str(e))
            return None
    
    async def _finalize_document_processing(
        self,
        document_id: str,
        session: AsyncSession
    ):
        """Finalize document processing with any additional indexing"""
        try:
            # Could add additional processing here like:
            # - Full-text search indexing
            # - Document classification refinement
            # - Quality scoring
            # - Relationship detection with other documents
            
            logger.debug("Document processing finalized", document_id=document_id)
            
        except Exception as e:
            logger.warning("Document finalization failed", document_id=document_id, error=str(e))
    
    async def get_processing_status(
        self,
        document_id: str,
        user_id: str,
        session: AsyncSession
    ) -> Optional[Dict[str, Any]]:
        """Get document processing status"""
        try:
            from sqlalchemy import select
            
            stmt = select(Document).where(
                Document.id == document_id,
                Document.user_id == user_id
            )
            
            result = await session.execute(stmt)
            document = result.scalar_one_or_none()
            
            if not document:
                return None
            
            return {
                'document_id': document_id,
                'status': document.status.value,
                'progress': document.processing_progress,
                'message': document.processing_message,
                'started_at': document.processing_started_at,
                'completed_at': document.processing_completed_at
            }
            
        except Exception as e:
            logger.error("Failed to get processing status", document_id=document_id, error=str(e))
            return None
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get processing statistics"""
        return self.processing_stats.copy()