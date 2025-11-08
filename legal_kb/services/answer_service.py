"""
Answer and Drafting Engine
- Defines a strict JSON contract for grounded answers
- Renders to Markdown and optionally .docx with footnote-style citations
"""
from __future__ import annotations

from dataclasses import dataclass, asdict
from typing import Dict, List, Optional


@dataclass
class CitedStatute:
    id: str
    section: Optional[str]
    quote: Optional[str]
    as_on: Optional[str]


@dataclass
class CitedCase:
    id: str
    citation: Optional[str]
    para: Optional[str]
    quote: Optional[str]


@dataclass
class AnalysisEntry:
    issue: str
    application: str


@dataclass
class DraftFields:
    parties: str
    court: str
    facts: str
    grounds: List[str]
    reliefs: List[str]
    prayer: str


@dataclass
class Draft:
    type: str  # anticipatory_bail | quash | written_submissions
    fields: DraftFields


@dataclass
class Confidence:
    score: float
    reasons: List[str]


@dataclass
class AnswerContract:
    short_answer: str
    statutes: List[CitedStatute]
    cases: List[CitedCase]
    analysis: List[AnalysisEntry]
    draft: Optional[Draft]
    confidence: Confidence
    warnings: List[str]


class RenderService:
    def to_markdown(self, answer: AnswerContract) -> str:
        s = []
        s.append("# Short Answer\n\n" + (answer.short_answer or "").strip())
        if answer.statutes:
            s.append("\n## Statutes\n")
            for st in answer.statutes:
                s.append(f"- {st.id} (as-on: {st.as_on or 'today'})\n")
                if st.quote:
                    s.append(f"  > {st.quote}\n")
        if answer.cases:
            s.append("\n## Cases\n")
            for c in answer.cases:
                meta = f" ({c.citation})" if c.citation else ""
                pin = f" at {c.para}" if c.para else ""
                s.append(f"- {c.id}{meta}{pin}\n")
                if c.quote:
                    s.append(f"  > {c.quote}\n")
        if answer.analysis:
            s.append("\n## Application to Facts\n")
            for a in answer.analysis:
                s.append(f"- **{a.issue}**: {a.application}\n")
        if answer.draft:
            s.append(f"\n## Draft ({answer.draft.type})\n")
            d = answer.draft.fields
            s.append(f"### Parties\n{d.parties}\n\n### Court\n{d.court}\n\n### Facts\n{d.facts}\n")
            if d.grounds:
                s.append("\n### Grounds\n" + "\n".join([f"- {g}" for g in d.grounds]))
            if d.reliefs:
                s.append("\n\n### Reliefs\n" + "\n".join([f"- {r}" for r in d.reliefs]))
            s.append("\n\n### Prayer\n" + d.prayer + "\n")
        s.append(f"\n## Confidence\nScore: {answer.confidence.score:.2f}\n")
        if answer.warnings:
            s.append("\n## Warnings\n" + "\n".join([f"- {w}" for w in answer.warnings]))
        return "\n".join(s).strip() + "\n"

    def to_docx(self, answer: AnswerContract, out_path: str) -> Optional[str]:
        try:
            from docx import Document  # type: ignore
        except Exception:
            return None
        doc = Document()
        doc.add_heading("Short Answer", level=1)
        doc.add_paragraph(answer.short_answer or "")
        if answer.statutes:
            doc.add_heading("Statutes", level=2)
            for st in answer.statutes:
                p = doc.add_paragraph(f"{st.id} (as-on: {st.as_on or 'today'})")
                if st.quote:
                    doc.add_paragraph(st.quote, style='Intense Quote')
        if answer.cases:
            doc.add_heading("Cases", level=2)
            for c in answer.cases:
                meta = f" ({c.citation})" if c.citation else ""
                pin = f" at {c.para}" if c.para else ""
                doc.add_paragraph(f"{c.id}{meta}{pin}")
                if c.quote:
                    doc.add_paragraph(c.quote, style='Intense Quote')
        if answer.analysis:
            doc.add_heading("Application to Facts", level=2)
            for a in answer.analysis:
                doc.add_paragraph(f"{a.issue}: {a.application}")
        if answer.draft:
            doc.add_heading(f"Draft ({answer.draft.type})", level=2)
            d = answer.draft.fields
            doc.add_heading("Parties", level=3); doc.add_paragraph(d.parties)
            doc.add_heading("Court", level=3); doc.add_paragraph(d.court)
            doc.add_heading("Facts", level=3); doc.add_paragraph(d.facts)
            if d.grounds:
                doc.add_heading("Grounds", level=3)
                for g in d.grounds:
                    doc.add_paragraph(g, style='List Bullet')
            if d.reliefs:
                doc.add_heading("Reliefs", level=3)
                for r in d.reliefs:
                    doc.add_paragraph(r, style='List Bullet')
            doc.add_heading("Prayer", level=3); doc.add_paragraph(d.prayer)
        doc.add_heading("Confidence", level=2)
        doc.add_paragraph(f"Score: {answer.confidence.score:.2f}")
        if answer.warnings:
            doc.add_heading("Warnings", level=2)
            for w in answer.warnings:
                doc.add_paragraph(w, style='List Bullet')
        doc.save(out_path)
        return out_path

